import json
import os
import matplotlib
matplotlib.use('Agg')  # Use a non-GUI backend before importing pyplot
import matplotlib.pyplot as plt
import pandas as pd
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_chroma import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader
from flask import Flask, request, jsonify, send_from_directory, send_file
from docx import Document
from docx.shared import Inches

chat_history = []

# Flask app setup
app = Flask(__name__)

# Load data from directories
prec_loader = DirectoryLoader('output/precipitation/')
temp_loader = DirectoryLoader('output/temperature/')
pres_loader = DirectoryLoader('output/pressure/')

prec_data = prec_loader.load()
temp_data = temp_loader.load()
pres_data = pres_loader.load()

# Combine all documents
all_documents = prec_data + temp_data + pres_data

# Split documents into chunks
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=50,
    length_function=len,
    is_separator_regex=False
)
split_documents = text_splitter.split_documents(all_documents)

# Load credentials
with open('secrets.json') as f:
    secrets = json.load(f)
secrets = secrets['gpt-4o']

embeddings = AzureOpenAIEmbeddings(
    azure_endpoint=secrets['endpoint'],
    api_key=secrets['api_key']
)

# Index configuration
index_metadata = {
    "hnsw:space": "cosine",
    "hnsw:M": 32,
    "hnsw:ef_construction": 128
}

# Check if DB exists, if not create and persist it
if not os.path.exists('db'):
    store = Chroma.from_documents(
        split_documents,
        embeddings,
        ids=[f"{doc.metadata.get('source', 'unknown-source')}-{i}" for i, doc in enumerate(split_documents)],
        persist_directory='db',
        collection_metadata=index_metadata
    )
    store.persist()
else:
    store = Chroma(persist_directory='db', embedding_function=embeddings)

# Retrieve settings
num_documents = len(store.get()["ids"])
k_value = min(num_documents, 40)

retriever = store.as_retriever(
    search_type="mmr",  # Improves document diversity
    search_kwargs={"k": k_value, "fetch_k": num_documents}
)

# Optimized prompt
prompt = PromptTemplate(
    template="""Consider {context} where:
    - Year: the year
    - Month: the month
    - Mean Precipitation (mm): average precipitation for the month
    - Max Precipitation (mm): maximum precipitation in a single day during the month
    - Min Precipitation (mm): minimum precipitation in a single day during the month
    - Mode Precipitation (mm): most frequent precipitation value for the month
    - Mean Temperature (°C): average monthly temperature
    - Max Temperature (°C): maximum monthly temperature
    - Min Temperature (°C): minimum monthly temperature
    - Mode Temperature (°C): most frequent monthly temperature value
    - Mean Pressure (mbar): average atmospheric pressure for the month
    - Max Pressure (mbar): highest atmospheric pressure recorded during the month
    - Min Pressure (mbar): lowest atmospheric pressure recorded during the month
    - Mode Pressure (mbar): most frequent atmospheric pressure value for the month

    If the question is about precipitation, use the relevant data.  
    If the question is about temperature, use the temperature-related values.  
    If the question is about atmospheric pressure, use the pressure-related values.

    If any calculations are required, explicitly show the calculations in a readable format:
        - Use plain text equations without LaTeX formatting.
        - If the equation is complex, provide a step-by-step breakdown.
        - Optionally, generate an image representation of the equation for clarity.

    Answer the following question: {question}""",
    input_variables=['context', 'question'],
)

# Configure AI model
llm = AzureChatOpenAI(
    azure_endpoint=secrets['endpoint'],
    api_key=secrets['api_key'],
    api_version="2023-12-01-preview",
    deployment_name=secrets['deployment_name'],
    model="model_finetuned_name"
)

llm_chain = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type='stuff',
    retriever=retriever,
    chain_type_kwargs={'prompt': prompt},
    return_source_documents=False
)

# Serve HTML front page
@app.route('/')
def home():
    return send_from_directory(os.getcwd(), 'index.html')

# Chat endpoint
@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.json.get('message')
    if not user_input:
        return jsonify({"error": "No message provided"}), 400

    try:
        response = llm_chain.invoke(user_input)
        response_text = response if isinstance(response, str) else response.get("result", "No response available")
        response_text = response_text.replace("**", "").replace("###", "")
        chat_history.append({"question": user_input, "answer": response_text})
        return jsonify({"response": response_text})
    except Exception as e:
        print(f"Error generating response: {e}")
        return jsonify({"error": "Error generating response"}), 500

# Download report endpoint
@app.route('/download', methods=['GET'])
def download_chat():
    if not chat_history:
        return jsonify({"error": "No conversation available"}), 400

    conversation_text = "\n".join([f"Q: {entry['question']}\nA: {entry['answer']}" for entry in chat_history])

    # Generate title
    title_prompt = f"Generate a title for this conversation (don't include the word 'title'):\n{conversation_text}"
    try:
        title_response = llm_chain.invoke(title_prompt)
        title = title_response if isinstance(title_response, str) else title_response.get("result", "Climate Report")
    except Exception as e:
        print(f"Error generating title: {e}")
        title = "Climate Report"

    # Generate summary
    summary_prompt = f"Imagine you are a meteorologist writing an abstract for this report (3-4 sentences):\n{conversation_text}"
    try:
        summary_response = llm_chain.invoke(summary_prompt)
        summary = summary_response if isinstance(summary_response, str) else summary_response.get("result", "Summary unavailable.")
    except Exception as e:
        print(f"Error generating summary: {e}")
        summary = "Summary unavailable."

    # Create Word document
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_heading("Abstract", 1)
    doc.add_paragraph(summary)
    doc.add_page_break()

    for entry in chat_history:
        doc.add_heading(entry["question"], level=1)
        doc.add_paragraph(entry["answer"])

    doc.add_page_break()
    doc.add_heading("Summary Table", level=1)

    # Prompt to generate summary table
    table_prompt = (
        "Generate a summary table of the data in CSV format. "
        "Each row must have two columns: "
        "the first column should be a keyword summarizing the question, "
        "the second column should contain the numerical data extracted from the answer. "
        "Do not include additional text beyond the table.\n"
        f"{conversation_text}"
    )
    try:
        table_response = llm_chain.invoke(table_prompt)
        table_text = table_response if isinstance(table_response, str) else table_response.get("result", "Table not available")
    except Exception as e:
        print(f"Error generating the table: {e}")
        table_text = "Table not available"

    # Create Word table
    if table_text and table_text != "Table not available":
        rows = [row.strip() for row in table_text.split("\n") if row.strip()]
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Question"
        header_cells[1].text = "Numerical Data"

        data_groups = {}
        for i, row in enumerate(rows):
            columns = [col.strip() for col in row.split(",")]
            if len(columns) >= 2:
                question, value_str = columns[0], columns[1]

                row_cells = table.add_row().cells
                row_cells[0].text = question
                row_cells[1].text = value_str

                try:
                    value = float(value_str)
                except ValueError:
                    print(f"Non-numeric value in row: {row}")
                    continue

                parts = question.strip().split(" ", 1)
                if parts[0].isdigit():
                    label = parts[0]
                    group_key = parts[1]
                else:
                    label = question.strip()
                    group_key = question.strip()

                if group_key not in data_groups:
                    data_groups[group_key] = []
                data_groups[group_key].append((label, value))
            else:
                print(f"[Row {i}] Invalid row: '{row}'")
    else:
        doc.add_paragraph("Table not available.")

    doc.add_heading("Data Charts", level=1)

    for group_key, values in data_groups.items():
        if not values:
            continue

        labels = [item[0] for item in values]
        data_values = [item[1] for item in values]

        plt.figure(figsize=(max(6, len(labels)), 4))

        if len(data_values) == 1:
            plt.text(0.5, 0.5, f"{data_values[0]}", fontsize=28, ha='center', va='center')
            plt.xlabel('Value')
            plt.axis('off')
        else:
            plt.bar(labels, data_values, color='mediumseagreen')
            plt.ylabel('Value')

        plt.title(group_key)
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()

        chart_path = f"{group_key.replace(' ', '_')}.png"
        plt.savefig(chart_path)
        plt.close()

        doc.add_paragraph(group_key)
        doc.add_picture(chart_path, width=Inches(6))

    # Save and return the Word document
    file_path = "report.docx"
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
